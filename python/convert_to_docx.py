#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将项目报告的 Markdown 文件转换为 Word (.docx) 格式"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import os
import sys

# Windows 控制台默认 GBK，统一改为 UTF-8 输出避免中文乱码
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_styled_paragraph(doc, text, style_name, font_name='Microsoft YaHei', font_size=None, bold=False, color=None, alignment=None):
    """添加带样式的段落"""
    p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_table_from_md(doc, lines, start_idx):
    """解析并添加Markdown表格到Word"""
    # 收集表格行
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    if not table_lines:
        return i

    # 跳过表头分隔行
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows_data.append(cells)

    # 过滤分隔行 (----)
    filtered_rows = []
    for row in rows_data:
        if not all(re.match(r'^[-:]+$', c) for c in row):
            filtered_rows.append(row)

    if not filtered_rows:
        return i

    # 创建Word表格
    ncols = max(len(r) for r in filtered_rows)
    table = doc.add_table(rows=len(filtered_rows), cols=ncols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(filtered_rows):
        row = table.rows[r_idx]
        # 补齐列
        while len(row_data) < ncols:
            row_data.append('')
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(9)

            # 表头加粗+蓝底
            if r_idx == 0:
                run.font.bold = True
                run.font.size = Pt(10)
                set_cell_shading(cell, '4472C4')
                run.font.color.rgb = RGBColor(255, 255, 255)
            # 偶数行灰底
            elif r_idx % 2 == 0:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()  # 表后空行
    return i

def convert_markdown_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8-sig') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_lines = []
    in_quote = False
    quote_lines = []
    list_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            if in_code_block:
                # 输出代码块
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(1)
                # 灰底
                pPr = p._element.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F5F5F5',
                    qn('w:val'): 'clear'
                })
                pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            # 一级标题
            title_text = stripped[2:].strip()
            p = add_styled_paragraph(doc, title_text, None, font_size=22, bold=True, color=(33, 150, 243), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()  # 空行
            i += 1
            continue

        if stripped.startswith('## '):
            title_text = stripped[3:].strip()
            add_styled_paragraph(doc, title_text, None, font_size=16, bold=True, color=(33, 150, 243))
            i += 1
            continue

        if stripped.startswith('### '):
            title_text = stripped[4:].strip()
            add_styled_paragraph(doc, title_text, None, font_size=13, bold=True, color=(66, 66, 66))
            i += 1
            continue

        if stripped.startswith('#### '):
            title_text = stripped[5:].strip()
            add_styled_paragraph(doc, title_text, None, font_size=11, bold=True, color=(66, 66, 66))
            i += 1
            continue

        # 引用块
        if stripped.startswith('> '):
            quote_lines.append(stripped[2:])
            in_quote = True
            i += 1
            # 检查下一行是否还是引用
            if i < len(lines) and lines[i].strip().startswith('> '):
                continue
            # 输出引用
            quote_text = '\n'.join(quote_lines)
            p = doc.add_paragraph()
            run = p.add_run(quote_text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.left_indent = Cm(1.5)
            # 左边框效果 - 用缩进+底色
            quote_lines = []
            in_quote = False
            continue

        # 水平线
        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # 无序列表
        if re.match(r'^- ', stripped):
            list_text = re.sub(r'^- ', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(list_text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+\. ', stripped):
            list_text = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            run = p.add_run(list_text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
            i += 1
            continue

        # 粗体文本行
        if stripped.startswith('**') and stripped.endswith('**'):
            bold_text = stripped[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(bold_text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
            run.font.bold = True
            i += 1
            continue

        # 普通段落
        if stripped:
            # 处理内联格式
            p = doc.add_paragraph()
            # 处理粗体 **text**
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(6)
        else:
            # 空行
            pass

        i += 1

    # 保存
    doc.save(docx_path)
    return True

if __name__ == '__main__':
    # 脚本位于 python/ 子目录，docs/ 在项目根目录，需上跳一级
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    md_path = os.path.join(project_dir, 'docs', 'project_report.md')
    docx_path = os.path.join(project_dir, 'docs', 'NovelMart电商经营分析项目报告.docx')

    print("正在转换项目报告 Markdown → Word...")
    convert_markdown_to_docx(md_path, docx_path)

    size_kb = os.path.getsize(docx_path) / 1024
    print(f"[OK] Word报告已生成: {docx_path}")
    print(f"[OK] 文件大小: {size_kb:.1f} KB")
